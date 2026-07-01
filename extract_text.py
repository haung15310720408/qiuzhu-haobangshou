#!/usr/bin/env python3
import json
import sys
from pathlib import Path


def trim(text, limit=200000):
    text = (text or "").strip()
    return text[:limit]


def read_pdf(path):
    import pdfplumber

    parts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages[:80]:
            parts.append(page.extract_text() or "")
    return "\n".join(parts)


def read_docx(path):
    import docx

    document = docx.Document(path)
    parts = [p.text for p in document.paragraphs if p.text]
    for table in document.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells]
            if any(values):
                parts.append("\t".join(values))
    return "\n".join(parts)


def read_xlsx(path):
    import openpyxl

    workbook = openpyxl.load_workbook(path, read_only=True, data_only=True)
    parts = []
    for sheet in workbook.worksheets[:10]:
        parts.append(f"# {sheet.title}")
        for idx, row in enumerate(sheet.iter_rows(values_only=True)):
            if idx >= 5000:
                break
            values = ["" if value is None else str(value) for value in row]
            if any(value.strip() for value in values):
                parts.append("\t".join(values))
    return "\n".join(parts)


def main():
    path = Path(sys.argv[1])
    ext = path.suffix.lower().lstrip(".")
    if ext == "pdf":
        text = read_pdf(path)
    elif ext == "docx":
        text = read_docx(path)
    elif ext in ("xlsx", "xlsm"):
        text = read_xlsx(path)
    else:
        text = path.read_text(encoding="utf-8", errors="ignore")
    print(json.dumps({"ok": True, "text": trim(text)}, ensure_ascii=False))


if __name__ == "__main__":
    try:
        main()
    except Exception as exc:
        print(json.dumps({"ok": False, "error": str(exc)}, ensure_ascii=False))
        sys.exit(1)
